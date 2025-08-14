import os
import datetime as dt
import tkinter as tk
from tkinter import filedialog, messagebox
import docx
from docx2pdf import convert


class InvoiceAutomation:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Invoice Automation")
        self.root.geometry('500x600')

        self.partner_label = tk.Label(self.root, text='Partner')
        self.partner_street_label = tk.Label(self.root, text='Partner Street')
        self.partner_ZIP_City_Country_label = tk.Label(
            self.root, text='Partner ZIP/City/Country')
        self.invoice_number_label = tk.Label(self.root, text='Invoice Number')
        self.service_decription_label = tk.Label(
            self.root, text='Service Description')
        self.service_amount_label = tk.Label(self.root, text='Service Amount')
        self.service_single_price_label = tk.Label(
            self.root, text='Service Single Price')
        self.payment_method_label = tk.Label(self.root, text='Payment Method')

        self.payment_methods = {
            'Main Bank': {
                'Recipient': 'Reliance Company',
                'Bank': 'SBI',
                'IBAN': 'IN123456789',
                'BIC': 'SBIIN123',
            },
            'Second Bank': {
                'Recipient': 'Reliance Company',
                'Bank': 'Axis Bank',
                'IBAN': 'IN123456789',
                'BIC': 'AIXIN123',
            },
            'Private Bank': {
                'Recipient': 'Reliance Company',
                'Bank': 'PNB',
                'IBAN': 'IN123456789',
                'BIC': 'PNBIIN123',
            }
        }

        self.partner_entry = tk.Entry(self.root)
        self.partner_street_entry = tk.Entry(self.root,)
        self.partner_ZIP_City_Country_entry = tk.Entry(self.root)
        self.invoice_number_entry = tk.Entry(self.root)
        self.service_decription_entry = tk.Entry(self.root)
        self.service_amount_entry = tk.Entry(self.root)
        self.service_single_price_entry = tk.Entry(self.root)

        self.payment_method = tk.StringVar(self.root)
        self.payment_method.set('Main Bank')

        self.payment_method_dropdown = tk.OptionMenu(
            self.root, self.payment_method, "Main Bank", "Second Bank", "Private Bank")

        self.create_button = tk.Button(
            self.root, text='Create Invoice', command=self.create_invoice)

        padding_options = {'fill': 'x', 'expand': True, 'padx': 10, 'pady': 4}

        self.partner_label.pack(padding_options)
        self.partner_entry.pack(padding_options)
        self.partner_street_label.pack(padding_options)
        self.partner_street_entry.pack(padding_options)
        self.partner_ZIP_City_Country_label.pack(padding_options)
        self.partner_ZIP_City_Country_entry.pack(padding_options)
        self.invoice_number_label.pack(padding_options)
        self.invoice_number_entry.pack(padding_options)
        self.service_decription_label.pack(padding_options)
        self.service_decription_entry.pack(padding_options)
        self.service_amount_label.pack(padding_options)
        self.service_amount_entry.pack(padding_options)
        self.service_single_price_label.pack(padding_options)
        self.service_single_price_entry.pack(padding_options)
        self.payment_method_label.pack(padding_options)
        self.payment_method_dropdown.pack(padding_options)
        self.create_button.pack(pady=20)

        self.root.mainloop()

    def paragraph_replace_text(self, paragraph, replacements):
        """
        Replaces text in a paragraph, handling placeholders split across runs.
        This is a more robust replacement function.
        """
        full_text = "".join(run.text for run in paragraph.runs)
        replaced_text = full_text

        for old, new in replacements.items():
            # Use str(new) to ensure the replacement value is a string
            replaced_text = replaced_text.replace(old, str(new))

        if full_text != replaced_text:
            # Clear all existing runs in the paragraph
            for run in paragraph.runs:
                r = run._element
                r.getparent().remove(r)
            # Add a new run with the fully replaced text
            paragraph.add_run(replaced_text)

    def create_invoice(self):
        script_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(script_dir, 'template.docx')

        try:
            doc = docx.Document(template_path)
        except docx.opc.exceptions.PackageNotFoundError:
            messagebox.showerror(
                title='Error', message=f"The 'template.docx' file was not found.\nPlease make sure it is in the same folder as the script:\n{script_dir}")
            return

        selected_payment_method = self.payment_methods[self.payment_method.get(
        )]

        try:
            if not all([self.partner_entry.get(), self.partner_street_entry.get(), self.invoice_number_entry.get(), self.service_amount_entry.get(), self.service_single_price_entry.get()]):
                messagebox.showerror(
                    title='Error', message='Please fill in all required fields.')
                return

            # --- CHANGE: Replaced '$' with '₹' for Rupee symbol ---
            replacements = {
                "[Date]": dt.datetime.today().strftime('%Y-%m-%d'),
                "[Partner]": self.partner_entry.get(),
                "[Partner Street]": self.partner_street_entry.get(),
                "[Partner ZIP_City_Country]": self.partner_ZIP_City_Country_entry.get(),
                "[Invoice_Number]": self.invoice_number_entry.get(),
                "[Service Description]": self.service_decription_entry.get(),
                "[Amount]": self.service_amount_entry.get(),
                "[Single Price]": f"₹{float(self.service_single_price_entry.get()):.2f}",
                "[Full Price]": f'₹{float(self.service_amount_entry.get()) * float(self.service_single_price_entry.get()):.2f}',
                "[Recipient]": selected_payment_method['Recipient'],
                "[Bank]": selected_payment_method['Bank'],
                "[IBAN]": selected_payment_method['IBAN'],
                "[BIC]": selected_payment_method['BIC'],
            }
        except ValueError:
            messagebox.showerror(
                title='Error', message='Invalid input for amount or price. Please use only numbers.')
            return

        for paragraph in doc.paragraphs:
            self.paragraph_replace_text(paragraph, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.paragraph_replace_text(paragraph, replacements)

        save_path = filedialog.asksaveasfilename(
            defaultextension='.pdf', filetypes=[('PDF documents', '*.pdf')])

        if not save_path:
            return

        try:
            filled_docx_path = os.path.join(
                script_dir, f"temp_{os.path.basename(save_path)}.docx")
            doc.save(filled_docx_path)
            convert(filled_docx_path, save_path)
            os.remove(filled_docx_path)
            messagebox.showinfo(
                'Success', f'Invoice created and saved successfully at:\n{save_path}')
        except Exception as e:
            messagebox.showerror(
                title='Error', message=f'Failed to create PDF. Please ensure Microsoft Word is installed and closed.\n\nDetails: {e}')


if __name__ == '__main__':
    InvoiceAutomation()
